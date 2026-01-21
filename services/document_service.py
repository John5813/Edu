import os
import re
import json
import logging
from typing import Dict
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK

logger = logging.getLogger(__name__)

class DocumentService:
    # existing methods remain; we add new helpers and update create_course_work

    def _get_course_work_texts(self, language: str) -> Dict[str, str]:
        # keep existing implementation if present; simplified fallback here
        if language == 'ru':
            return {
                'course_work': 'КУРСОВАЯ РАБОТА',
                'faculty': 'факультет',
                'topic': 'Тема',
                'prepared_by': 'Выполнил(а)',
                'accepted_by': 'Принял(а)',
                'city': 'Ташкент',
                'contents': 'СОДЕРЖАНИЕ',
                'introduction': 'ВВЕДЕНИЕ',
                'chapter': 'ГЛАВА',
                'conclusion': 'ЗАКЛЮЧЕНИЕ',
                'references': 'СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ',
                'intro_points': []
            }
        elif language == 'en':
            return {
                'course_work': 'COURSE WORK',
                'faculty': 'faculty',
                'topic': 'Topic',
                'prepared_by': 'Prepared by',
                'accepted_by': 'Accepted by',
                'city': 'Tashkent',
                'contents': 'CONTENTS',
                'introduction': 'INTRODUCTION',
                'chapter': 'CHAPTER',
                'conclusion': 'CONCLUSION',
                'references': 'REFERENCES',
                'intro_points': []
            }
        else:
            return {
                'course_work': 'KURS ISHI',
                'faculty': 'fakulteti',
                'topic': 'Mavzu',
                'prepared_by': 'Bajardi',
                'accepted_by': 'Qabul qildi',
                'city': 'Toshkent',
                'contents': 'MUNDARIJA',
                'introduction': 'KIRISH',
                'chapter': "BO'LIM",
                'conclusion': 'XULOSA',
                'references': 'FOYDALANILGAN ADABIYOTLAR',
                'intro_points': []
            }

    def extract_text_from_docx(self, file_path: str) -> str:
        """Read all paragraphs from a docx and return plain text for editing."""
        doc = DocxDocument(file_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        return "\n\n".join(parts)

    async def create_doc_from_edited_text(self, topic: str, edited_text: str, author_name: str, language: str) -> str:
        """
        Create a docx file from a single edited_text blob (simple structure).
        Returns the new file path.
        """
        safe = re.sub(r'\W+', '_', topic.lower())[:50]
        docx_fn = f"course_work_edited_{safe}.docx"
        out_dir = os.path.join("output", "course_works")
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, docx_fn)

        texts = self._get_course_work_texts(language)

        doc = DocxDocument()
        # Title page
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texts['course_work'] + "\n\n")
        run.bold = True
        run.font.size = Pt(20)
        doc.add_paragraph(f"{texts['topic']}: {topic}")
        doc.add_paragraph(f"{texts['prepared_by']}: {author_name}")
        doc.add_page_break()

        # Put edited text (simple) into the document
        for block in edited_text.split("\n\n"):
            para = doc.add_paragraph(block)
            para.paragraph_format.line_spacing = 1.5
            if para.runs:
                run = para.runs[0]
            else:
                run = para.add_run(block)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        doc.save(out_path)
        return out_path

    async def create_course_work(self, topic: str, content: dict, author_name: str, language: str) -> str:
        """
        Create a .docx file for the course work using structured content produced by AI.
        This method implements inline superscript markers and a Footnotes section at the end.
        """
        doc = DocxDocument()
        safe = re.sub(r'\W+', '_', topic.lower())[:50]
        docx_fn = f"course_work_{safe}.docx"
        out_dir = os.path.join("output", "course_works")
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, docx_fn)

        texts = self._get_course_work_texts(language)

        # Title page
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(texts['course_work'] + "\n\n")
        run.font.size = Pt(20)
        run.bold = True
        doc.add_paragraph(f"{texts['topic']}: {topic}")
        doc.add_paragraph(f"{texts['prepared_by']}: {author_name}")
        doc.add_page_break()

        # Table of contents (simple)
        doc.add_heading(texts['contents'], level=1)
        for ch_idx, ch in enumerate(content.get("chapters", []), start=1):
            toc_p = doc.add_paragraph()
            toc_p.add_run(f"{texts.get('chapter','Chapter')} {ch_idx}: {ch.get('title','')}")
            for sub_idx, sub in enumerate(ch.get("subsections", []), start=1):
                toc_sub = doc.add_paragraph(style='List Number')
                toc_sub.add_run(f"{sub_idx}. {sub.get('title','')}")
        doc.add_page_break()

        # Collect global footnotes
        global_footnotes = {}
        encountered_ids = []

        def add_paragraph_with_footnotes(par_doc, text):
            parts = re.split(r'(\[\d+\])', text)
            for part in parts:
                if re.fullmatch(r'\[\d+\]', part):
                    num = re.sub(r'\D', '', part)
                    run = par_doc.add_run(num)
                    font = run.font
                    font.superscript = True
                    font.size = Pt(10)
                else:
                    par_doc.add_run(part)

        # Write content
        if content.get("title"):
            doc.add_heading(content.get("title"), level=1)

        for ch_idx, ch in enumerate(content.get("chapters", []), start=1):
            doc.add_heading(f"{texts.get('chapter','CHAPTER')} {ch_idx}. {ch.get('title','')}", level=2)
            for sub_idx, sub in enumerate(ch.get("subsections", []), start=1):
                doc.add_heading(f"{sub_idx}. {sub.get('title','')}", level=3)
                p = doc.add_paragraph()
                add_paragraph_with_footnotes(p, sub.get("text",""))

                for fn in sub.get("footnotes", []):
                    try:
                        fid = int(fn.get("id"))
                        if fid not in global_footnotes:
                            global_footnotes[fid] = fn.get("text","")
                            if fid not in encountered_ids:
                                encountered_ids.append(fid)
                    except Exception:
                        continue

            doc.add_page_break()

        # top-level footnotes
        for fn in content.get("footnotes", []):
            try:
                fid = int(fn.get("id"))
                if fid not in global_footnotes:
                    global_footnotes[fid] = fn.get("text","")
                    if fid not in encountered_ids:
                        encountered_ids.append(fid)
            except Exception:
                continue

        # Append Footnotes section
        if encountered_ids:
            doc.add_heading("Footnotes", level=2)
            for fid in sorted(encountered_ids):
                text = global_footnotes.get(fid, "")
                p = doc.add_paragraph()
                run_num = p.add_run(f"{fid}. ")
                run_num.bold = True
                p.add_run(text)

        # References
        refs = content.get("references", [])
        if refs:
            doc.add_page_break()
            doc.add_heading(texts.get('references','REFERENCES'), level=2)
            for r in refs:
                doc.add_paragraph(r, style='List Bullet')

        doc.save(out_path)
        return out_path