"""Utilities for reading uploaded files that contain numbered tests."""

import re
from pathlib import Path


_QUESTION_START = re.compile(r"(?m)^[ \t]*(\d{1,4})[ \t]*[.)\-:][ \t]+")


def extract_file_text(file_path: str) -> str:
    """Extract text from a DOCX or PDF test file."""
    suffix = Path(file_path).suffix.lower()
    if suffix == ".docx":
        from docx import Document

        doc = Document(file_path)
        parts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells).strip()
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    if suffix == ".pdf":
        import fitz

        with fitz.open(file_path) as pdf:
            return "\n".join(page.get_text() for page in pdf)

    raise ValueError("Only DOCX and PDF files are supported")


def extract_numbered_tests(file_path: str) -> list[dict]:
    """Return numbered question blocks from a DOCX/PDF file."""
    text = extract_file_text(file_path).replace("\r\n", "\n").replace("\r", "\n")
    matches = list(_QUESTION_START.finditer(text))
    tests: list[dict] = []

    for index, match in enumerate(matches):
        block_end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        block = text[match.start():block_end].strip()
        if len(block) < 8:
            continue
        tests.append({"number": int(match.group(1)), "text": block})

    if not tests:
        raise ValueError("No numbered tests found")
    return tests